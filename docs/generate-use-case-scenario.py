#!/usr/bin/env python3
"""Generate use case scenario DOCX for ArsipDesa thesis."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

OUTPUT = Path(__file__).parent / "use-case-scenario.docx"

SCENARIOS = [
    {
        "table_no": "3.6",
        "name": "Login",
        "id": "UC01",
        "importance": "High",
        "primary_actor": "Admin, Sekdes, Kades",
        "use_case_type": "Supporting Use Case",
        "stakeholder": "Pengguna ingin mengakses sistem sesuai hak akses perannya.",
        "brief": "Use case ini menjelaskan bagaimana pengguna melakukan autentikasi sebelum mengakses sistem ArsipDesa.",
        "trigger": "Pengguna ingin mengakses dashboard sistem",
        "type": "Primary",
        "relationship": "Association",
        "include": "-",
        "extend": "Logout (UC10)",
        "generalization": "-",
        "normal_flow": [
            "Pengguna membuka halaman login sistem.",
            "Pengguna memasukkan username dan password, lalu menekan tombol masuk.",
            "Sistem memverifikasi kredensial pengguna.",
            "Sistem membuat sesi pengguna dan mengarahkan ke halaman dashboard sesuai peran (Admin, Sekdes, atau Kades).",
        ],
        "exceptional": "Jika username atau password salah, sistem menampilkan pesan gagal login dan pengguna tetap berada di halaman login.",
    },
    {
        "table_no": "3.7",
        "name": "Mengelola Surat Masuk & Keluar",
        "id": "UC02",
        "importance": "High",
        "primary_actor": "Admin",
        "use_case_type": "Main Use Case",
        "stakeholder": "Admin ingin mencatat dan mengelola data surat masuk serta surat keluar desa secara terpusat.",
        "brief": "Use case ini menjelaskan bagaimana Admin menambah, mengubah, menghapus, dan melihat data surat masuk serta surat keluar beserta dokumen pendukungnya.",
        "trigger": "Admin memilih menu Surat Masuk atau Surat Keluar",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Admin membuka menu Surat Masuk atau Surat Keluar.",
            "Sistem menampilkan daftar surat yang belum diarsipkan.",
            "Admin memilih menambah, mengubah, menghapus, atau melihat detail surat.",
            "Admin mengisi atau memperbarui data surat (nomor, tanggal, pengirim/tujuan, perihal, dan dokumen jika ada), lalu menyimpan.",
            "Sistem memvalidasi data, menyimpan perubahan, dan menampilkan pesan berhasil.",
        ],
        "exceptional": "Jika data tidak lengkap atau format file dokumen tidak sesuai (bukan PDF/DOC/DOCX), sistem menampilkan pesan kesalahan validasi dan data tidak disimpan.",
    },
    {
        "table_no": "3.8",
        "name": "Mengelola Arsip Surat",
        "id": "UC03",
        "importance": "High",
        "primary_actor": "Admin",
        "use_case_type": "Main Use Case",
        "stakeholder": "Admin ingin mengarsipkan surat yang telah selesai diproses agar tersimpan sebagai arsip resmi.",
        "brief": "Use case ini menjelaskan bagaimana Admin mengarsipkan surat masuk atau surat keluar yang telah didisposisikan, serta membatalkan arsip jika diperlukan.",
        "trigger": "Admin membuka detail surat yang siap diarsipkan",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Admin membuka detail surat masuk atau surat keluar dari menu terkait.",
            "Admin memilih aksi mengarsipkan surat.",
            "Sistem memeriksa apakah surat telah memenuhi syarat arsip (status didisposisikan untuk surat masuk).",
            "Sistem mengubah status surat menjadi diarsipkan dan mencatat waktu pengarsipan.",
            "Sistem menampilkan pesan berhasil dan surat dipindahkan ke daftar arsip.",
        ],
        "exceptional": "Jika surat belum didisposisikan, sistem menolak pengarsipan dan menampilkan pesan bahwa surat harus didisposisikan terlebih dahulu.",
    },
    {
        "table_no": "3.9",
        "name": "Mengelola User",
        "id": "UC04",
        "importance": "Medium",
        "primary_actor": "Admin",
        "use_case_type": "Main Use Case",
        "stakeholder": "Admin ingin mengelola akun pengguna sistem agar hak akses sesuai jabatan di kantor desa.",
        "brief": "Use case ini menjelaskan bagaimana Admin menambah, mengubah, dan menghapus data pengguna (Admin, Sekdes, Kades) pada sistem.",
        "trigger": "Admin memilih menu Manajemen User",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Admin membuka menu Manajemen User.",
            "Sistem menampilkan daftar pengguna terdaftar.",
            "Admin memilih menambah, mengubah, atau menghapus pengguna.",
            "Admin mengisi atau memperbarui data pengguna (nama, username, email, password, dan peran), lalu menyimpan.",
            "Sistem memvalidasi data, menyimpan perubahan, dan menampilkan pesan berhasil.",
        ],
        "exceptional": "Jika Admin mencoba menghapus akun sendiri atau mengubah peran akun sendiri, sistem menolak aksi tersebut dan menampilkan pesan kesalahan.",
    },
    {
        "table_no": "3.10",
        "name": "Review Surat Masuk",
        "id": "UC05",
        "importance": "High",
        "primary_actor": "Sekdes",
        "use_case_type": "Main Use Case",
        "stakeholder": "Sekretaris Desa ingin meninjau kelengkapan surat masuk dan menetapkan tingkat surat sesuai prosedur administrasi.",
        "brief": "Use case ini menjelaskan bagaimana Sekdes melakukan review terhadap surat masuk berstatus draft dan menetapkan tingkat surat (biasa atau penting).",
        "trigger": "Sekdes membuka detail surat masuk yang belum direview",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Sekdes membuka menu Surat Masuk dan memilih surat berstatus draft.",
            "Sekdes meninjau kelengkapan data surat pada halaman detail.",
            "Sekdes memilih tingkat surat (biasa atau penting) dan mengonfirmasi review.",
            "Sistem memperbarui tingkat surat, mengubah status menjadi terverifikasi, dan mencatat waktu serta pelaku review.",
            "Sistem menampilkan pesan berhasil.",
        ],
        "exceptional": "Jika surat sudah direview atau bukan berstatus draft, sistem menolak aksi review dan menampilkan pesan bahwa surat tidak dapat direview.",
    },
    {
        "table_no": "3.11",
        "name": "Verifikasi Surat Masuk",
        "id": "UC06",
        "importance": "High",
        "primary_actor": "Kades",
        "use_case_type": "Main Use Case",
        "stakeholder": "Kepala Desa ingin memverifikasi surat masuk penting sebelum diproses lebih lanjut.",
        "brief": "Use case ini menjelaskan bagaimana Kades melakukan verifikasi terhadap surat masuk berkategori penting yang telah direview Sekdes.",
        "trigger": "Kades membuka detail surat masuk penting yang menunggu verifikasi",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Kades membuka menu Surat Masuk dan memilih surat penting berstatus terverifikasi yang belum diverifikasi.",
            "Kades meninjau isi surat pada halaman detail.",
            "Kades mengonfirmasi verifikasi surat.",
            "Sistem mencatat waktu dan pelaku verifikasi Kades.",
            "Sistem menampilkan pesan berhasil.",
        ],
        "exceptional": "Jika surat bukan tingkat penting, belum direview Sekdes, atau sudah diverifikasi, sistem menolak verifikasi dan menampilkan pesan kesalahan.",
    },
    {
        "table_no": "3.12",
        "name": "Pemberian Disposisi",
        "id": "UC07",
        "importance": "High",
        "primary_actor": "Sekdes, Kades",
        "use_case_type": "Main Use Case",
        "stakeholder": "Sekdes dan Kades ingin memberikan disposisi surat kepada jabatan terkait agar surat segera ditindaklanjuti.",
        "brief": "Use case ini menjelaskan bagaimana Sekdes atau Kades memberikan disposisi surat masuk dengan menentukan jabatan tujuan dan catatan disposisi.",
        "trigger": "Sekdes atau Kades membuka surat masuk yang memenuhi syarat disposisi",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Pengguna membuka menu Disposisi atau halaman detail surat masuk.",
            "Sistem menampilkan surat yang memenuhi syarat disposisi sesuai peran pengguna.",
            "Pengguna memilih jabatan tujuan disposisi, mengisi catatan, dan menyimpan disposisi.",
            "Sistem menyimpan data disposisi, memperbarui status surat menjadi didisposisikan, dan menampilkan pesan berhasil.",
        ],
        "exceptional": "Jika surat belum memenuhi syarat (Sekdes: surat biasa belum terverifikasi; Kades: surat penting belum diverifikasi), sistem menolak pembuatan disposisi dan menampilkan pesan kesalahan.",
    },
    {
        "table_no": "3.13",
        "name": "Pencarian Arsip Surat",
        "id": "UC08",
        "importance": "Medium",
        "primary_actor": "Admin, Sekdes, Kades",
        "use_case_type": "Main Use Case",
        "stakeholder": "Pengguna ingin menemukan arsip surat dengan cepat berdasarkan nomor surat.",
        "brief": "Use case ini menjelaskan bagaimana pengguna melakukan pencarian arsip surat menggunakan prefix nomor surat dengan algoritma Binary Search, serta filter jenis dan rentang waktu.",
        "trigger": "Pengguna memilih menu Arsip Surat dan memasukkan kata kunci pencarian",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Pengguna membuka menu Arsip Surat.",
            "Pengguna memasukkan prefix nomor surat dan/atau memilih filter jenis surat (masuk/keluar) serta rentang waktu.",
            "Sistem melakukan pencarian arsip menggunakan algoritma Binary Search berdasarkan nomor surat.",
            "Sistem menampilkan daftar arsip surat yang sesuai kriteria pencarian.",
            "Pengguna dapat membuka detail arsip surat yang dipilih.",
        ],
        "exceptional": "Jika tidak ada arsip yang cocok dengan kata kunci, sistem menampilkan daftar kosong dan pesan bahwa data tidak ditemukan.",
    },
    {
        "table_no": "3.14",
        "name": "Rekapitulasi & Laporan Arsip Surat",
        "id": "UC09",
        "importance": "Medium",
        "primary_actor": "Admin, Sekdes, Kades",
        "use_case_type": "Main Use Case",
        "stakeholder": "Pengguna ingin memperoleh ringkasan data surat dan arsip untuk keperluan monitoring serta pelaporan administrasi desa.",
        "brief": "Use case ini menjelaskan bagaimana pengguna melihat rekapitulasi data surat masuk, surat keluar, arsip, dan disposisi, serta mengunduh laporan dalam format PDF.",
        "trigger": "Pengguna memilih menu Laporan",
        "type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal_flow": [
            "Pengguna membuka menu Laporan.",
            "Pengguna memilih rentang waktu laporan (semua waktu, 7 hari, 30 hari, 90 hari, atau 1 tahun).",
            "Sistem menampilkan ringkasan statistik surat masuk, surat keluar, arsip, disposisi, dan grafik tren.",
            "Pengguna dapat memilih mengunduh laporan dalam format PDF.",
            "Sistem menghasilkan dan mengunduh file laporan PDF.",
        ],
        "exceptional": "Jika tidak terdapat data pada periode yang dipilih, sistem tetap menampilkan laporan dengan nilai nol pada setiap ringkasan statistik.",
    },
    {
        "table_no": "3.15",
        "name": "Logout",
        "id": "UC10",
        "importance": "High",
        "primary_actor": "Admin, Sekdes, Kades",
        "use_case_type": "Supporting Use Case",
        "stakeholder": "Pengguna ingin keluar dari sistem secara aman setelah selesai menggunakan aplikasi.",
        "brief": "Use case ini menjelaskan bagaimana pengguna mengakhiri sesi dan keluar dari sistem ArsipDesa.",
        "trigger": "Pengguna memilih menu Logout pada header aplikasi",
        "type": "Secondary",
        "relationship": "Association",
        "include": "-",
        "extend": "Login (UC01)",
        "generalization": "-",
        "normal_flow": [
            "Pengguna memilih opsi Logout pada menu profil di header.",
            "Sistem mengakhiri sesi pengguna dan menghapus data autentikasi.",
            "Sistem mengarahkan pengguna ke halaman login.",
        ],
        "exceptional": "Jika sesi pengguna sudah berakhir, sistem tetap mengarahkan pengguna ke halaman login.",
    },
]


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge, val in kwargs.items():
        tag = f"w:{edge}"
        element = OxmlElement(tag)
        for key, value in val.items():
            element.set(qn(f"w:{key}"), str(value))
        tc_pr.append(element)


def apply_table_borders(table) -> None:
    border = {"val": "single", "sz": "4", "color": "000000", "space": "0"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def merge_row(table, row_idx: int, label: str, value: str, label_span: int = 1, value_span: int = 5) -> None:
    row = table.rows[row_idx]
    cells = row.cells
    set_cell_text(cells[0], label, bold=True)
    if label_span > 1:
        cells[0].merge(cells[label_span - 1])
    start = label_span
    set_cell_text(cells[start], value)
    if value_span > 1:
        cells[start].merge(cells[start + value_span - 1])


def add_scenario_table(doc: Document, scenario: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Tabel {scenario['table_no']} Use Case Scenario {scenario['name']}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = doc.add_table(rows=9, cols=6)
    table.autofit = False
    apply_table_borders(table)

    widths = [Cm(3.2), Cm(3.5), Cm(2.0), Cm(2.5), Cm(2.8), Cm(3.5)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    # Row 0: Use Case Name | value | ID | value | Importance | value
    r0 = table.rows[0].cells
    set_cell_text(r0[0], "Use Case Name:", bold=True)
    set_cell_text(r0[1], scenario["name"])
    set_cell_text(r0[2], "ID:", bold=True)
    set_cell_text(r0[3], scenario["id"])
    set_cell_text(r0[4], "Importance:", bold=True)
    set_cell_text(r0[5], scenario["importance"])

    # Row 1: Primary Actor | value | Use Case Type | value (span last 2 cols for type value)
    r1 = table.rows[1].cells
    set_cell_text(r1[0], "Primary Actor:", bold=True)
    set_cell_text(r1[1], scenario["primary_actor"])
    set_cell_text(r1[2], "Use Case Type:", bold=True)
    r1[3].merge(r1[5])
    set_cell_text(r1[3], scenario["use_case_type"])

    # Row 2: Stakeholder (full width value)
    merge_row(table, 2, "Stakeholder and Interest:", scenario["stakeholder"], 1, 5)

    # Row 3: Brief Description
    merge_row(table, 3, "Brief Description:", scenario["brief"], 1, 5)

    # Row 4: Trigger | Type | Relationship
    r4 = table.rows[4].cells
    set_cell_text(r4[0], "Trigger:", bold=True)
    set_cell_text(r4[1], scenario["trigger"])
    set_cell_text(r4[2], "Type:", bold=True)
    set_cell_text(r4[3], scenario["type"])
    set_cell_text(r4[4], "Relationship:", bold=True)
    set_cell_text(r4[5], scenario["relationship"])

    # Row 5: Include | Extend | Generalization
    r5 = table.rows[5].cells
    set_cell_text(r5[0], "Include:", bold=True)
    set_cell_text(r5[1], scenario["include"])
    set_cell_text(r5[2], "Extend:", bold=True)
    set_cell_text(r5[3], scenario["extend"])
    set_cell_text(r5[4], "Generalization/Inheritance:", bold=True)
    set_cell_text(r5[5], scenario["generalization"])

    # Row 6: Normal Flow label
    merge_row(table, 6, "Normal Flow of Events:", "", 1, 5)

    # Row 7: normal flow steps
    flow_text = "\n".join(
        f"{i}. {step}" for i, step in enumerate(scenario["normal_flow"], start=1)
    )
    merge_row(table, 7, "", flow_text, 1, 5)

    # Row 8: Exceptional Flows
    merge_row(table, 8, "Exceptional Flows:", scenario["exceptional"], 1, 5)

    doc.add_paragraph()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for scenario in SCENARIOS:
        add_scenario_table(doc, scenario)

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT} ({len(SCENARIOS)} scenarios)")


if __name__ == "__main__":
    main()
