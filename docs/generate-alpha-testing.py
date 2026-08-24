#!/usr/bin/env python3
"""Generate alpha testing tables (DOCX) for ArsipDesa thesis."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).parent / "pengujian-alpha.docx"

# Tabel 3.21–3.31: pengujian alpha per fitur (setelah data dictionary 3.16–3.20)
SECTIONS = [
    {
        "table_no": "3.21",
        "title": "Tabel pengujian alpha login",
        "cases": [
            (
                "Menekan tombol *login* dengan *username* dan *password* benar",
                "Sistem menampilkan halaman dashboard sesuai hak akses pengguna (Admin, Sekdes, atau Kades)",
            ),
            (
                "Menekan tombol *login* dengan *username* atau *password* salah",
                "Sistem menampilkan pesan 'Username atau password salah' dan tetap pada halaman login",
            ),
            (
                "Menekan tombol *login* tanpa mengisi *username* dan *password*",
                "Sistem menampilkan pesan validasi bahwa *username* dan *password* wajib diisi",
            ),
        ],
    },
    {
        "table_no": "3.22",
        "title": "Tabel pengujian alpha mengelola surat masuk",
        "cases": [
            (
                "Admin membuka menu Surat Masuk",
                "Sistem menampilkan daftar surat masuk yang belum diarsipkan",
            ),
            (
                "Admin mengisi form tambah surat masuk dengan data lengkap lalu menyimpan",
                "Sistem menyimpan data dan menampilkan pesan 'Surat Masuk berhasil ditambahkan'",
            ),
            (
                "Admin menyimpan form surat masuk dengan data wajib kosong",
                "Sistem menampilkan pesan validasi pada field yang belum diisi dan data tidak disimpan",
            ),
            (
                "Admin mengubah data surat masuk lalu menyimpan",
                "Sistem memperbarui data dan menampilkan pesan 'Surat Masuk berhasil diperbarui'",
            ),
            (
                "Admin menghapus data surat masuk",
                "Sistem menghapus data dan menampilkan pesan 'Surat Masuk berhasil dihapus'",
            ),
            (
                "Admin membuka detail surat masuk",
                "Sistem menampilkan detail surat masuk beserta file dokumen jika tersedia",
            ),
        ],
    },
    {
        "table_no": "3.23",
        "title": "Tabel pengujian alpha mengelola surat keluar",
        "cases": [
            (
                "Admin membuka menu Surat Keluar",
                "Sistem menampilkan daftar surat keluar yang belum diarsipkan",
            ),
            (
                "Admin mengisi form tambah surat keluar dengan data lengkap lalu menyimpan",
                "Sistem menyimpan data dan menampilkan pesan 'Surat Keluar berhasil ditambahkan'",
            ),
            (
                "Admin menyimpan form surat keluar dengan data wajib kosong",
                "Sistem menampilkan pesan validasi pada field yang belum diisi dan data tidak disimpan",
            ),
            (
                "Admin mengubah data surat keluar lalu menyimpan",
                "Sistem memperbarui data dan menampilkan pesan 'Surat Keluar berhasil diperbarui'",
            ),
            (
                "Admin menghapus data surat keluar",
                "Sistem menghapus data dan menampilkan pesan 'Surat Keluar berhasil dihapus'",
            ),
            (
                "Admin membuka detail surat keluar",
                "Sistem menampilkan detail surat keluar beserta file dokumen jika tersedia",
            ),
        ],
    },
    {
        "table_no": "3.24",
        "title": "Tabel pengujian alpha mengelola arsip surat",
        "cases": [
            (
                "Admin mengarsipkan surat masuk yang berstatus didisposisikan",
                "Sistem mengarsipkan surat dan menampilkan pesan 'Surat masuk berhasil diarsipkan'",
            ),
            (
                "Admin mengarsipkan surat masuk yang belum didisposisikan",
                "Sistem menampilkan pesan 'Surat harus didisposisikan terlebih dahulu sebelum diarsipkan'",
            ),
            (
                "Admin mengarsipkan surat keluar",
                "Sistem mengarsipkan surat dan menampilkan pesan 'Surat keluar berhasil diarsipkan'",
            ),
            (
                "Admin membatalkan arsip surat masuk",
                "Sistem mengembalikan surat ke daftar aktif dan menampilkan pesan 'Surat masuk dikembalikan ke daftar aktif'",
            ),
            (
                "Admin membatalkan arsip surat keluar",
                "Sistem mengembalikan surat ke daftar aktif dan menampilkan pesan 'Surat keluar dikembalikan ke daftar aktif'",
            ),
            (
                "Admin / Sekdes / Kades membuka menu Arsip Surat",
                "Sistem menampilkan daftar surat masuk dan surat keluar yang sudah diarsipkan",
            ),
        ],
    },
    {
        "table_no": "3.25",
        "title": "Tabel pengujian alpha mengelola user",
        "cases": [
            (
                "Admin membuka menu Manajemen User",
                "Sistem menampilkan daftar pengguna sistem",
            ),
            (
                "Admin menambah pengguna baru dengan data lengkap lalu menyimpan",
                "Sistem menyimpan data dan menampilkan pesan 'Pengguna berhasil ditambahkan'",
            ),
            (
                "Admin menambah pengguna dengan data wajib kosong",
                "Sistem menampilkan pesan validasi dan data pengguna tidak disimpan",
            ),
            (
                "Admin mengubah data pengguna lain lalu menyimpan",
                "Sistem memperbarui data dan menampilkan pesan 'Pengguna berhasil diperbarui'",
            ),
            (
                "Admin mencoba mengubah *role* akun sendiri",
                "Sistem menampilkan pesan 'Anda tidak dapat mengubah peran akun sendiri'",
            ),
            (
                "Admin menghapus pengguna lain",
                "Sistem menghapus data dan menampilkan pesan 'Pengguna berhasil dihapus'",
            ),
            (
                "Admin mencoba menghapus akun sendiri",
                "Sistem menampilkan pesan 'Tidak dapat menghapus akun sendiri'",
            ),
        ],
    },
    {
        "table_no": "3.26",
        "title": "Tabel pengujian alpha review surat masuk",
        "cases": [
            (
                "Sekdes membuka detail surat masuk berstatus *draft*",
                "Sistem menampilkan detail surat dan opsi review tingkat surat",
            ),
            (
                "Sekdes mereview surat masuk dengan menetapkan tingkat biasa atau penting",
                "Sistem memperbarui status menjadi terverifikasi dan menampilkan pesan 'Surat berhasil direview dan tingkat ditetapkan'",
            ),
            (
                "Sekdes mencoba mereview surat yang bukan berstatus *draft*",
                "Sistem menolak aksi review karena surat tidak memenuhi syarat review",
            ),
            (
                "Pengguna selain Sekdes mencoba mengakses aksi review Sekdes",
                "Sistem menolak akses (unauthorized / tidak menampilkan aksi review)",
            ),
        ],
    },
    {
        "table_no": "3.27",
        "title": "Tabel pengujian alpha verifikasi surat masuk",
        "cases": [
            (
                "Kades membuka detail surat masuk tingkat penting yang sudah direview Sekdes",
                "Sistem menampilkan detail surat dan opsi verifikasi Kades",
            ),
            (
                "Kades menekan tombol verifikasi pada surat penting yang memenuhi syarat",
                "Sistem mencatat verifikasi dan menampilkan pesan 'Surat penting berhasil diverifikasi'",
            ),
            (
                "Kades mencoba memverifikasi surat tingkat biasa",
                "Sistem menolak verifikasi karena surat tidak memenuhi syarat verifikasi Kades",
            ),
            (
                "Kades mencoba memverifikasi surat yang belum direview Sekdes",
                "Sistem menolak verifikasi karena surat belum berstatus terverifikasi / belum direview",
            ),
            (
                "Pengguna selain Kades mencoba mengakses aksi verifikasi Kades",
                "Sistem menolak akses (unauthorized / tidak menampilkan aksi verifikasi)",
            ),
        ],
    },
    {
        "table_no": "3.28",
        "title": "Tabel pengujian alpha pemberian disposisi",
        "cases": [
            (
                "Sekdes / Kades membuka menu Disposisi atau form disposisi dari detail surat",
                "Sistem menampilkan form disposisi beserta daftar surat yang eligible",
            ),
            (
                "Sekdes membuat disposisi untuk surat tingkat biasa yang sudah terverifikasi",
                "Sistem menyimpan disposisi, mengubah status surat menjadi didisposisikan, dan menampilkan pesan 'Disposisi berhasil dikirim'",
            ),
            (
                "Kades membuat disposisi untuk surat tingkat penting yang sudah diverifikasi Kades",
                "Sistem menyimpan disposisi, mengubah status surat menjadi didisposisikan, dan menampilkan pesan 'Disposisi berhasil dikirim'",
            ),
            (
                "Pengguna membuat disposisi tanpa mengisi catatan / jabatan tujuan",
                "Sistem menampilkan pesan validasi bahwa field wajib harus diisi",
            ),
            (
                "Pengguna membuat disposisi pada surat yang belum memenuhi syarat",
                "Sistem menampilkan pesan 'Surat belum memenuhi syarat untuk dibuatkan disposisi'",
            ),
            (
                "Admin mencoba membuat disposisi",
                "Sistem menolak akses karena Admin tidak berwenang membuat disposisi",
            ),
        ],
    },
    {
        "table_no": "3.29",
        "title": "Tabel pengujian alpha pencarian arsip surat",
        "cases": [
            (
                "Pengguna membuka menu Arsip Surat",
                "Sistem menampilkan halaman arsip surat masuk dan surat keluar",
            ),
            (
                "Pengguna mencari arsip dengan prefix nomor surat yang ada di database",
                "Sistem menampilkan daftar arsip yang nomor suratnya cocok dengan prefix pencarian",
            ),
            (
                "Pengguna mencari arsip dengan nomor surat yang tidak ditemukan",
                "Sistem menampilkan daftar kosong / pesan bahwa data tidak ditemukan",
            ),
            (
                "Pengguna memfilter arsip berdasarkan jenis surat (masuk / keluar)",
                "Sistem menampilkan hanya arsip sesuai filter jenis yang dipilih",
            ),
            (
                "Pengguna membuka detail arsip surat",
                "Sistem menampilkan detail lengkap arsip surat yang dipilih",
            ),
        ],
    },
    {
        "table_no": "3.30",
        "title": "Tabel pengujian alpha rekapitulasi dan laporan",
        "cases": [
            (
                "Pengguna membuka menu Laporan",
                "Sistem menampilkan halaman rekapitulasi statistik surat dan arsip",
            ),
            (
                "Pengguna memilih rentang waktu laporan lalu menampilkan data",
                "Sistem menampilkan statistik dan grafik sesuai rentang waktu yang dipilih",
            ),
            (
                "Pengguna menekan tombol unduh laporan PDF",
                "Sistem mengunduh file PDF laporan surat sesuai filter yang aktif",
            ),
            (
                "Pengguna melihat laporan tanpa memilih filter rentang waktu",
                "Sistem menampilkan data laporan dengan rentang default (seluruh data / all)",
            ),
        ],
    },
    {
        "table_no": "3.31",
        "title": "Tabel pengujian alpha logout",
        "cases": [
            (
                "Pengguna menekan tombol *logout* saat sesi aktif",
                "Sistem mengakhiri sesi pengguna dan mengarahkan ke halaman login",
            ),
            (
                "Pengguna mencoba mengakses halaman dashboard setelah *logout*",
                "Sistem mengarahkan kembali ke halaman login karena sesi sudah berakhir",
            ),
        ],
    },
]


def set_run_font(run, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_mixed_text(paragraph, text: str, *, bold: bool = False, size: int = 11) -> None:
    """Render text with *word* as italic (like skripsi format)."""
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, bold=bold, italic=True, size=size)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, bold=bold, italic=False, size=size)


def set_cell_border(cell) -> None:
    border = {"val": "single", "sz": "4", "color": "000000", "space": "0"}
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        for key, value in border.items():
            element.set(qn(f"w:{key}"), str(value))
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_vmerge(cell, restart: bool = False) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    vmerge = OxmlElement("w:vMerge")
    vmerge.set(qn("w:val"), "restart" if restart else "continue")
    tc_pr.append(vmerge)


def clear_cell(cell) -> None:
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)


def write_header_cell(cell, text: str, *, center: bool = True) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, bold=True, size=11)
    set_cell_shading(cell, "BDD7EE")
    set_cell_border(cell)


def write_body_cell(cell, text: str = "", *, center: bool = False) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    if text:
        add_mixed_text(p, text, size=11)
    set_cell_border(cell)


def add_alpha_table(doc: Document, section: dict) -> None:
    # Judul: Tabel 3. 21 Tabel pengujian alpha login
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    add_mixed_text(title, f"Tabel {section['table_no']} {section['title']}", bold=True)

    cases = section["cases"]
    # 2 header rows + N case rows, 4 columns
    table = doc.add_table(rows=2 + len(cases), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(6.5), Cm(7.0), Cm(2.0), Cm(2.5)]

    # Header row 0
    r0 = table.rows[0].cells
    for i, w in enumerate(widths):
        r0[i].width = w

    write_header_cell(r0[0], "Test Case")
    write_header_cell(r0[1], "Yang diharapkan")
    write_header_cell(r0[2], "Hasil Pengujian")
    write_header_cell(r0[3], "")
    # Merge Hasil Pengujian across last two columns
    r0[2].merge(r0[3])
    # Re-apply text after merge
    write_header_cell(r0[2], "Hasil Pengujian")

    # Header row 1 — Sesuai / Tidak Sesuai
    r1 = table.rows[1].cells
    for i, w in enumerate(widths):
        r1[i].width = w

    write_header_cell(r1[0], "")
    write_header_cell(r1[1], "")
    write_header_cell(r1[2], "Sesuai")
    write_header_cell(r1[3], "Tidak Sesuai")

    # Vertical merge Test Case & Yang diharapkan header cells
    set_cell_vmerge(r0[0], restart=True)
    set_cell_vmerge(r1[0], restart=False)
    set_cell_vmerge(r0[1], restart=True)
    set_cell_vmerge(r1[1], restart=False)

    # Body
    for idx, (case, expected) in enumerate(cases):
        row = table.rows[2 + idx].cells
        for i, w in enumerate(widths):
            row[i].width = w
        write_body_cell(row[0], case)
        write_body_cell(row[1], expected)
        write_body_cell(row[2], "", center=True)
        write_body_cell(row[3], "", center=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(14)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    for spec in SECTIONS:
        add_alpha_table(doc, spec)

    doc.save(OUTPUT)
    total_cases = sum(len(s["cases"]) for s in SECTIONS)
    print(f"Generated: {OUTPUT} ({len(SECTIONS)} tables, {total_cases} test cases)")


if __name__ == "__main__":
    main()
