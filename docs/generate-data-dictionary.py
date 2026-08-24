#!/usr/bin/env python3
"""Generate database data dictionary tables (DOCX) for ArsipDesa thesis."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).parent / "tabel-data-dictionary.docx"

# Format mengikuti contoh skripsi: Tabel 3.16 dst.
# Field & tipe disesuaikan dengan migrasi/model aplikasi ArsipDesa.
TABLES = [
    {
        "table_no": "3.16",
        "title": "Tabel Data Users",
        "db_table": "users",
        "fields": [
            ("id", "bigint(20)", "Primary key, auto-increment, ID unik pengguna"),
            ("name", "varchar(255)", "Nama lengkap pengguna"),
            ("username", "varchar(255)", "Nama pengguna (unik) yang digunakan untuk login"),
            ("email", "varchar(255)", "Alamat email pengguna (unik)"),
            ("email_verified_at", "timestamp", "Tanggal dan waktu email diverifikasi (nullable)"),
            ("password", "varchar(255)", "Password terenkripsi sebagai kredensial login"),
            ("role", "enum('admin','sekdes','kades')", "Role / tipe akun untuk otorisasi akses sistem"),
            ("remember_token", "varchar(100)", "Token untuk fitur remember me (nullable)"),
            ("created_at", "timestamp", "Tanggal dan waktu data dibuat"),
            ("updated_at", "timestamp", "Tanggal dan waktu data diperbarui"),
        ],
    },
    {
        "table_no": "3.17",
        "title": "Tabel Data Surat Masuk",
        "db_table": "surat_masuk",
        "fields": [
            ("id", "bigint(20)", "Primary key, auto-increment, ID unik surat masuk"),
            ("no_surat", "varchar(255)", "Nomor surat masuk (unik)"),
            ("tanggal_terima", "date", "Tanggal surat diterima"),
            ("tanggal_surat", "date", "Tanggal tertera pada surat (nullable)"),
            ("pengirim", "varchar(255)", "Nama instansi atau pihak pengirim surat"),
            ("perihal", "varchar(255)", "Perihal / subjek surat"),
            ("catatan", "text", "Catatan tambahan terkait surat (nullable)"),
            ("status", "varchar(255)", "Status alur surat: draft, terverifikasi, didisposisikan, diarsipkan"),
            ("tingkat", "varchar(16)", "Tingkat surat: biasa atau penting (nullable)"),
            ("tujuan", "varchar(255)", "Tujuan / penerima internal surat"),
            ("file", "varchar(255)", "Path file dokumen surat (nullable)"),
            ("diarsipkan_at", "timestamp", "Tanggal dan waktu surat diarsipkan (nullable)"),
            ("verified_sekdes_at", "timestamp", "Tanggal dan waktu direview oleh Sekdes (nullable)"),
            ("verified_sekdes_by", "bigint(20)", "Foreign key ke users.id — pengguna Sekdes yang mereview (nullable)"),
            ("verified_kades_at", "timestamp", "Tanggal dan waktu diverifikasi oleh Kades (nullable)"),
            ("verified_kades_by", "bigint(20)", "Foreign key ke users.id — pengguna Kades yang memverifikasi (nullable)"),
            ("created_at", "timestamp", "Tanggal dan waktu data dibuat"),
            ("updated_at", "timestamp", "Tanggal dan waktu data diperbarui"),
        ],
    },
    {
        "table_no": "3.18",
        "title": "Tabel Data Surat Keluar",
        "db_table": "surat_keluar",
        "fields": [
            ("id", "bigint(20)", "Primary key, auto-increment, ID unik surat keluar"),
            ("surat_masuk_id", "bigint(20)", "Foreign key ke surat_masuk.id — surat masuk terkait (nullable)"),
            ("no_surat", "varchar(255)", "Nomor surat keluar (unik)"),
            ("tanggal_kirim", "date", "Tanggal surat dikirim"),
            ("tujuan", "varchar(255)", "Tujuan / penerima surat keluar"),
            ("perihal", "varchar(255)", "Perihal / subjek surat"),
            ("catatan", "text", "Catatan tambahan terkait surat (nullable)"),
            ("status", "enum('draft','terkirim')", "Status surat keluar: draft atau terkirim"),
            ("file", "varchar(255)", "Path file dokumen surat"),
            ("diarsipkan_at", "timestamp", "Tanggal dan waktu surat diarsipkan (nullable)"),
            ("created_at", "timestamp", "Tanggal dan waktu data dibuat"),
            ("updated_at", "timestamp", "Tanggal dan waktu data diperbarui"),
        ],
    },
    {
        "table_no": "3.19",
        "title": "Tabel Data Disposisi",
        "db_table": "disposisi",
        "fields": [
            ("id", "bigint(20)", "Primary key, auto-increment, ID unik disposisi"),
            ("surat_masuk_id", "bigint(20)", "Foreign key ke surat_masuk.id — surat yang didisposisikan"),
            ("user_id", "bigint(20)", "Foreign key ke users.id — pengguna (Sekdes/Kades) yang membuat disposisi"),
            ("jabatan_tujuan_id", "bigint(20)", "Foreign key ke jabatan_tujuan_disposisi.id — jabatan tujuan (nullable)"),
            ("dari_jabatan", "varchar(64)", "Jabatan pemberi disposisi, misalnya Sekretaris Desa / Kepala Desa (nullable)"),
            ("kepada", "varchar(255)", "Nama jabatan tujuan disposisi (teks)"),
            ("catatan", "text", "Isi catatan / instruksi disposisi"),
            ("tanggal", "date", "Tanggal disposisi dibuat"),
            ("created_at", "timestamp", "Tanggal dan waktu data dibuat"),
            ("updated_at", "timestamp", "Tanggal dan waktu data diperbarui"),
        ],
    },
    {
        "table_no": "3.20",
        "title": "Tabel Data Jabatan Tujuan Disposisi",
        "db_table": "jabatan_tujuan_disposisi",
        "fields": [
            ("id", "bigint(20)", "Primary key, auto-increment, ID unik jabatan tujuan"),
            ("nama_jabatan", "varchar(255)", "Nama jabatan tujuan disposisi (unik)"),
            ("is_active", "boolean / tinyint(1)", "Status aktif jabatan (true = dapat dipilih)"),
            ("sort_order", "smallint unsigned", "Urutan tampilan jabatan pada form"),
            ("created_at", "timestamp", "Tanggal dan waktu data dibuat"),
            ("updated_at", "timestamp", "Tanggal dan waktu data diperbarui"),
        ],
    },
]


def set_run_font(run, *, bold: bool = False, size: int = 11) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_border(cell) -> None:
    border = {"val": "single", "sz": "4", "color": "000000", "space": "0"}
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        for key, value in border.items():
            element.set(qn(f"w:{key}"), str(value))
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    center: bool = False,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=11)


def add_data_table(doc: Document, spec: dict) -> None:
    # Judul seperti contoh: "Tabel 3. 16 Tabel Data Users"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(f"Tabel {spec['table_no']} {spec['title']}")
    set_run_font(run, bold=True, size=11)

    fields = spec["fields"]
    table = doc.add_table(rows=1 + len(fields), cols=4)
    table.autofit = False

    widths = [Cm(1.5), Cm(4.5), Cm(4.5), Cm(6.5)]
    headers = ["No", "Nama Field", "Tipe Data", "Keterangan"]

    # Header row — biru muda seperti contoh
    header_row = table.rows[0]
    for i, (header, width) in enumerate(zip(headers, widths)):
        cell = header_row.cells[i]
        cell.width = width
        set_cell_text(cell, header, bold=True, center=True)
        set_cell_shading(cell, "BDD7EE")
        set_cell_border(cell)

    # Body rows
    for idx, (name, tipe, ket) in enumerate(fields, start=1):
        row = table.rows[idx]
        values = [str(idx), name, tipe, ket]
        centers = [True, False, False, False]
        for i, (value, width, center) in enumerate(zip(values, widths, centers)):
            cell = row.cells[i]
            cell.width = width
            set_cell_text(cell, value, center=center)
            set_cell_border(cell)

    # Spasi antar tabel
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    for spec in TABLES:
        add_data_table(doc, spec)

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT} ({len(TABLES)} tables)")


if __name__ == "__main__":
    main()
